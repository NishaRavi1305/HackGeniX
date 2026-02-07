"""
Document processing service for resumes and job descriptions.

Handles PDF/DOCX text extraction and LLM-based structured parsing.
"""
import io
import re
import logging
import hashlib
import json
from typing import List, Dict, Any, Optional
from pathlib import Path

from pdfplumber import open as open_pdf
from docx import Document as DocxDocument

from src.models.documents import (
    ParsedResume,
    ParsedJobDescription,
    ParsedEntity,
    ContactInfo,
    Education,
    Experience,
    Project,
    Research,
)

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Document processing service for resumes and job descriptions.

    Uses LLM for structured extraction from document text.
    """

    def __init__(self):
        """Initialize the document processor."""
        logger.info("DocumentProcessor initialized (LLM-only mode)")

    # =========================================================================
    # Text Extraction Methods
    # =========================================================================

    async def extract_text_from_pdf(self, pdf_bytes: bytes) -> str:
        """
        Extract text content from a PDF file.

        Tries pdfplumber first (fast, works for normal PDFs). If the PDF is
        image-based or uses non-standard fonts (common with LaTeX/Overleaf
        exports), falls back to OCR via Tesseract.

        Args:
            pdf_bytes: PDF file content as bytes

        Returns:
            Extracted text content
        """
        # --- Attempt 1: pdfplumber (text-based PDFs) ---
        text_parts = []
        try:
            with open_pdf(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")

        text = "\n".join(text_parts).strip()
        if len(text) >= 50:
            return text

        # --- Attempt 2: OCR via PyMuPDF + Tesseract ---
        logger.info("pdfplumber returned little/no text, falling back to OCR")
        try:
            import fitz  # PyMuPDF
            import pytesseract
            from PIL import Image

            pytesseract.pytesseract.tesseract_cmd = (
                r"C:\Program Files\Tesseract-OCR\tesseract.exe"
            )

            ocr_parts = []
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            for page_num, page in enumerate(doc):
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                page_text = pytesseract.image_to_string(img)
                if page_text and page_text.strip():
                    ocr_parts.append(page_text)
                logger.debug(f"OCR page {page_num}: {len(page_text)} chars")
            doc.close()

            ocr_text = "\n".join(ocr_parts).strip()
            if ocr_text:
                logger.info(f"OCR extracted {len(ocr_text)} chars from PDF")
                return ocr_text

        except ImportError as e:
            logger.warning(f"OCR dependencies not available: {e}")
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")

        # If both methods fail, return whatever we got (possibly empty)
        if text:
            return text
        raise ValueError(
            "Could not extract text from PDF. The file may be image-based "
            "and OCR dependencies (pymupdf, pytesseract, Tesseract) are required."
        )

    async def extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """
        Extract text content from a DOCX file.

        Args:
            docx_bytes: DOCX file content as bytes

        Returns:
            Extracted text content
        """
        try:
            doc = DocxDocument(io.BytesIO(docx_bytes))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise

    async def extract_text(self, file_bytes: bytes, content_type: str) -> str:
        """
        Extract text from a document based on its content type.

        Args:
            file_bytes: File content as bytes
            content_type: MIME type of the file

        Returns:
            Extracted text content
        """
        if content_type == "application/pdf":
            return await self.extract_text_from_pdf(file_bytes)
        elif content_type in [
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]:
            return await self.extract_text_from_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported content type: {content_type}")

    # =========================================================================
    # LLM-Based Parsing Methods
    # =========================================================================

    async def parse_resume_with_llm(
        self,
        text: str,
        use_cache: bool = True,
    ) -> ParsedResume:
        """
        Parse a resume using LLM for robust extraction.

        This method uses an LLM to extract structured information from resume text,
        handling diverse formats and layouts that rule-based parsing may miss.

        Args:
            text: Extracted resume text
            use_cache: Whether to use cached results if available

        Returns:
            ParsedResume with extracted information
        """
        # Generate cache key from text hash
        text_hash = hashlib.sha256(text.encode()).hexdigest()[:16]
        cache_key = f"resume_{text_hash}"

        # Check cache
        if use_cache:
            cached = self._get_from_cache(cache_key)
            if cached:
                logger.info(f"Using cached resume parse: {cache_key}")
                return cached

        if not text or len(text.strip()) < 50:
            logger.warning("Resume text too short, returning empty parse")
            return ParsedResume(raw_text=text)

        try:
            from src.providers.llm import get_llm_provider
            from src.providers.llm.base import GenerationConfig, user_message, system_message
            from src.services.prompts import RESUME_EXTRACTION_PROMPT

            llm = await get_llm_provider()

            # Truncate text if too long (keep first 8000 chars for context window)
            resume_text = text[:8000] if len(text) > 8000 else text

            prompt = RESUME_EXTRACTION_PROMPT.format(resume_text=resume_text)

            messages = [
                system_message(
                    "You are a JSON formatter. Your ONLY job is to map raw resume text "
                    "into structured JSON key-value pairs. Copy every word verbatim. "
                    "Do not summarize, interpret, or rephrase anything."
                ),
                user_message(prompt),
            ]

            logger.info("Parsing resume with LLM...")
            response = await llm.generate(
                messages,
                GenerationConfig(max_tokens=4096, temperature=0.1)
            )

            # Parse JSON response
            data = self._parse_llm_json_response(response.content)

            if not data:
                logger.warning("LLM returned invalid JSON for resume, returning minimal parse")
                return ParsedResume(raw_text=text)

            # Convert to ParsedResume
            parsed = self._json_to_parsed_resume(data, text)

            # Safety-net: merge technologies from projects into skills
            parsed = self._enrich_skills(parsed)

            # Cache the result
            if use_cache:
                self._save_to_cache(cache_key, parsed)

            logger.info(f"LLM resume parse complete: {parsed.contact.name if parsed.contact else 'Unknown'}")
            return parsed

        except Exception as e:
            logger.warning(f"LLM resume parsing failed: {e}, returning minimal parse")
            return ParsedResume(raw_text=text)

    async def parse_job_description_with_llm(
        self,
        text: str,
        use_cache: bool = True,
    ) -> ParsedJobDescription:
        """
        Parse a job description using LLM for robust extraction.

        Args:
            text: Job description text
            use_cache: Whether to use cached results if available

        Returns:
            ParsedJobDescription with extracted information
        """
        # Generate cache key from text hash
        text_hash = hashlib.sha256(text.encode()).hexdigest()[:16]
        cache_key = f"jd_{text_hash}"

        # Check cache
        if use_cache:
            cached = self._get_from_cache(cache_key)
            if cached:
                logger.info(f"Using cached JD parse: {cache_key}")
                return cached

        if not text or len(text.strip()) < 50:
            logger.warning("JD text too short, returning empty parse")
            return ParsedJobDescription(raw_text=text)

        try:
            from src.providers.llm import get_llm_provider
            from src.providers.llm.base import GenerationConfig, user_message, system_message
            from src.services.prompts import JD_EXTRACTION_PROMPT

            llm = await get_llm_provider()

            # Truncate if too long
            jd_text = text[:8000] if len(text) > 8000 else text

            prompt = JD_EXTRACTION_PROMPT.format(jd_text=jd_text)

            messages = [
                system_message("You are an expert job description parser. Extract requirements accurately."),
                user_message(prompt),
            ]

            logger.info("Parsing job description with LLM...")
            response = await llm.generate(
                messages,
                GenerationConfig(max_tokens=2048, temperature=0.1)
            )

            # Parse JSON response
            data = self._parse_llm_json_response(response.content)

            if not data:
                logger.warning("LLM returned invalid JSON for JD, returning minimal parse")
                return ParsedJobDescription(raw_text=text)

            # Convert to ParsedJobDescription
            parsed = self._json_to_parsed_jd(data, text)

            # Cache the result
            if use_cache:
                self._save_to_cache(cache_key, parsed)

            logger.info(f"LLM JD parse complete: {parsed.title or 'Unknown'}")
            return parsed

        except Exception as e:
            logger.warning(f"LLM JD parsing failed: {e}, returning minimal parse")
            return ParsedJobDescription(raw_text=text)

    # =========================================================================
    # JSON Parsing Helpers
    # =========================================================================

    def _parse_llm_json_response(self, content: str) -> Optional[Dict[str, Any]]:
        """Parse JSON from LLM response, handling common formatting issues."""
        content = content.strip()

        # Handle markdown code blocks
        if "```json" in content:
            start = content.find("```json") + 7
            end = content.find("```", start)
            content = content[start:end].strip()
        elif "```" in content:
            start = content.find("```") + 3
            end = content.find("```", start)
            content = content[start:end].strip()

        # Find JSON object bounds
        if not content.startswith("{"):
            start = content.find("{")
            if start != -1:
                content = content[start:]

        if not content.endswith("}"):
            end = content.rfind("}")
            if end != -1:
                content = content[:end + 1]

        try:
            return json.loads(content)
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse LLM JSON: {e}")
            logger.debug(f"Content was: {content[:500]}")
            return None

    @staticmethod
    def _get_first(*values) -> Optional[str]:
        """Return the first non-None, non-empty value from the list."""
        for v in values:
            if v is not None and (not isinstance(v, str) or v.strip()):
                return v
        return None

    @staticmethod
    def _enrich_skills(parsed: ParsedResume) -> ParsedResume:
        """Merge skills from all sections into the top-level skills list.

        Collects technologies/skills from experience entries, projects
        (tech_stack + skills), and research entries, then deduplicates
        and merges into parsed.skills.
        """
        existing = set(s.lower().strip() for s in parsed.skills)
        extra: list[str] = []

        def _add(item: str) -> None:
            key = item.lower().strip()
            if key and key not in existing:
                extra.append(item.strip())
                existing.add(key)

        for exp in parsed.experience:
            for s in exp.skills:
                _add(s)
        for proj in parsed.projects:
            for s in proj.tech_stack:
                _add(s)
            for s in proj.skills:
                _add(s)
        for res in parsed.research:
            for s in res.skills:
                _add(s)

        if extra:
            parsed.skills.extend(extra)
            logger.debug(f"Enriched skills with {len(extra)} items: {extra}")

        return parsed

    def _json_to_parsed_resume(self, data: Dict[str, Any], raw_text: str) -> ParsedResume:
        """Convert LLM JSON output to ParsedResume model.

        Handles field name variations that small LLMs (e.g. qwen2.5:3b) may produce.
        """
        _first = self._get_first

        def _str_list(val: Any) -> List[str]:
            """Safely coerce a value to a list of strings."""
            if isinstance(val, list):
                return [s for s in val if isinstance(s, str)]
            return []

        # Extract contact info -- the LLM may nest it under various keys
        contact_data = (
            data.get("contact")
            or data.get("contact_info")
            or data.get("personal_info")
            or data.get("personal_details")
            or {}
        )
        if not isinstance(contact_data, dict):
            contact_data = {}

        contact = ContactInfo(
            name=_first(
                contact_data.get("name"),
                contact_data.get("full_name"),
                contact_data.get("candidate_name"),
                data.get("name"),
                data.get("full_name"),
            ),
            email=_first(
                contact_data.get("email"),
                contact_data.get("email_address"),
            ),
            phone=_first(
                contact_data.get("phone"),
                contact_data.get("phone_number"),
                contact_data.get("mobile"),
                contact_data.get("telephone"),
            ),
            linkedin=_first(
                contact_data.get("linkedin"),
                contact_data.get("linkedin_url"),
                contact_data.get("linkedin_profile"),
            ),
            github=_first(
                contact_data.get("github"),
                contact_data.get("github_url"),
                contact_data.get("github_profile"),
            ),
            location=_first(
                contact_data.get("location"),
                contact_data.get("address"),
                contact_data.get("city"),
            ),
        )

        # Extract experience
        experience_list = data.get("experience") or data.get("work_experience") or []
        experience = []
        for exp_data in experience_list:
            if isinstance(exp_data, dict):
                experience.append(Experience(
                    company=_first(
                        exp_data.get("company"),
                        exp_data.get("organization"),
                        exp_data.get("employer"),
                    ),
                    title=_first(
                        exp_data.get("title"),
                        exp_data.get("job_title"),
                        exp_data.get("role"),
                        exp_data.get("position"),
                    ),
                    location=exp_data.get("location"),
                    start_date=exp_data.get("start_date"),
                    end_date=exp_data.get("end_date"),
                    description=exp_data.get("description"),
                    highlights=_str_list(exp_data.get("highlights")),
                    skills=_str_list(exp_data.get("skills")),
                    impact=_str_list(exp_data.get("impact")),
                ))

        # Extract education
        education_list = data.get("education") or []
        education = []
        for edu_data in education_list:
            if isinstance(edu_data, dict):
                education.append(Education(
                    institution=_first(
                        edu_data.get("institution"),
                        edu_data.get("university"),
                        edu_data.get("school"),
                        edu_data.get("college"),
                    ),
                    degree=_first(
                        edu_data.get("degree"),
                        edu_data.get("qualification"),
                    ),
                    field=_first(
                        edu_data.get("field"),
                        edu_data.get("major"),
                        edu_data.get("specialization"),
                    ),
                    start_date=edu_data.get("start_date"),
                    end_date=edu_data.get("end_date"),
                    gpa=edu_data.get("gpa"),
                ))

        # Extract skills -- may be nested under various keys
        skills = (
            data.get("skills")
            or data.get("technical_skills")
            or data.get("skill_set")
            or []
        )
        skills = _str_list(skills)

        # Extract certifications
        certifications = _str_list(data.get("certifications"))

        # Extract projects
        projects_list = data.get("projects") or data.get("personal_projects") or []
        projects = []
        for proj_data in projects_list:
            if isinstance(proj_data, dict):
                projects.append(Project(
                    name=_first(
                        proj_data.get("name"),
                        proj_data.get("title"),
                        proj_data.get("project_name"),
                    ),
                    description=proj_data.get("description"),
                    tech_stack=_str_list(
                        proj_data.get("tech_stack")
                        or proj_data.get("technologies")
                    ),
                    highlights=_str_list(proj_data.get("highlights")),
                    skills=_str_list(proj_data.get("skills")),
                    impact=_str_list(proj_data.get("impact")),
                    url=proj_data.get("url"),
                ))

        # Extract research
        research_list = data.get("research") or data.get("publications") or []
        research = []
        for res_data in research_list:
            if isinstance(res_data, dict):
                research.append(Research(
                    title=_first(
                        res_data.get("title"),
                        res_data.get("name"),
                    ),
                    venue=_first(
                        res_data.get("venue"),
                        res_data.get("journal"),
                        res_data.get("conference"),
                        res_data.get("publication"),
                    ),
                    status=res_data.get("status"),
                    highlights=_str_list(res_data.get("highlights")),
                    skills=_str_list(res_data.get("skills")),
                    impact=_str_list(res_data.get("impact")),
                ))

        # Extract soft skills and areas of interest
        soft_skills = _str_list(data.get("soft_skills"))
        areas_of_interest = _str_list(
            data.get("areas_of_interest")
            or data.get("interests")
            or data.get("domains")
        )

        # Extra sections — anything not already consumed
        known_keys = {
            "contact", "contact_info", "personal_info", "personal_details",
            "summary", "education", "experience", "work_experience",
            "skills", "technical_skills", "skill_set",
            "certifications", "projects", "personal_projects",
            "research", "publications",
            "soft_skills", "areas_of_interest", "interests", "domains",
            "name", "full_name", "extraction_confidence",
        }
        extra_sections = {
            k: v for k, v in data.items()
            if k not in known_keys and not k.startswith("_")
        }

        return ParsedResume(
            contact=contact,
            summary=data.get("summary"),
            skills=skills,
            experience=experience,
            education=education,
            certifications=certifications,
            projects=projects,
            research=research,
            soft_skills=soft_skills,
            areas_of_interest=areas_of_interest,
            extra_sections=extra_sections,
            raw_text=raw_text,
        )

    def _json_to_parsed_jd(self, data: Dict[str, Any], raw_text: str) -> ParsedJobDescription:
        """Convert LLM JSON output to ParsedJobDescription model."""
        return ParsedJobDescription(
            title=data.get("title"),
            company=data.get("company"),
            location=data.get("location"),
            employment_type=data.get("employment_type"),
            experience_level=data.get("experience_level"),
            experience_years_min=data.get("experience_years_min"),
            experience_years_max=data.get("experience_years_max"),
            salary_min=data.get("salary_min"),
            salary_max=data.get("salary_max"),
            salary_currency=data.get("salary_currency"),
            required_skills=data.get("required_skills", []),
            preferred_skills=data.get("preferred_skills", []),
            responsibilities=data.get("responsibilities", []),
            qualifications=data.get("qualifications", []),
            raw_text=raw_text,
        )

    # =========================================================================
    # Caching Methods
    # =========================================================================

    _cache_dir: Optional[Path] = None

    @classmethod
    def _get_cache_dir(cls) -> Path:
        """Get or create the cache directory."""
        if cls._cache_dir is None:
            cache_path = Path(__file__).parent.parent.parent / ".cache" / "parsed_documents"
            cache_path.mkdir(parents=True, exist_ok=True)
            cls._cache_dir = cache_path
        return cls._cache_dir

    def _get_from_cache(self, cache_key: str) -> Optional[Any]:
        """Get a cached parsed document."""
        cache_file = self._get_cache_dir() / f"{cache_key}.json"

        if not cache_file.exists():
            return None

        try:
            with open(cache_file, "r", encoding="utf-8") as f:
                data = json.load(f)

            # Reconstruct the appropriate model
            if cache_key.startswith("resume_"):
                return self._json_to_parsed_resume(data, data.get("_raw_text", ""))
            elif cache_key.startswith("jd_"):
                return self._json_to_parsed_jd(data, data.get("_raw_text", ""))

            return None

        except Exception as e:
            logger.warning(f"Failed to load from cache: {e}")
            return None

    def _save_to_cache(self, cache_key: str, parsed: Any) -> None:
        """Save a parsed document to cache."""
        cache_file = self._get_cache_dir() / f"{cache_key}.json"

        try:
            # Convert model to dict for caching
            if isinstance(parsed, ParsedResume):
                data = {
                    "contact": {
                        "name": parsed.contact.name if parsed.contact else None,
                        "email": parsed.contact.email if parsed.contact else None,
                        "phone": parsed.contact.phone if parsed.contact else None,
                        "linkedin": parsed.contact.linkedin if parsed.contact else None,
                        "github": parsed.contact.github if parsed.contact else None,
                        "location": parsed.contact.location if parsed.contact else None,
                    },
                    "summary": parsed.summary,
                    "skills": parsed.skills,
                    "experience": [
                        {
                            "company": exp.company,
                            "title": exp.title,
                            "location": exp.location,
                            "start_date": exp.start_date,
                            "end_date": exp.end_date,
                            "description": exp.description,
                            "highlights": exp.highlights,
                            "skills": exp.skills,
                            "impact": exp.impact,
                        }
                        for exp in parsed.experience
                    ],
                    "education": [
                        {
                            "institution": edu.institution,
                            "degree": edu.degree,
                            "field": edu.field,
                            "start_date": edu.start_date,
                            "end_date": edu.end_date,
                            "gpa": edu.gpa,
                        }
                        for edu in parsed.education
                    ],
                    "certifications": parsed.certifications,
                    "projects": [
                        {
                            "name": proj.name,
                            "description": proj.description,
                            "tech_stack": proj.tech_stack,
                            "highlights": proj.highlights,
                            "skills": proj.skills,
                            "impact": proj.impact,
                            "url": proj.url,
                        }
                        for proj in parsed.projects
                    ],
                    "research": [
                        {
                            "title": res.title,
                            "venue": res.venue,
                            "status": res.status,
                            "highlights": res.highlights,
                            "skills": res.skills,
                            "impact": res.impact,
                        }
                        for res in parsed.research
                    ],
                    "soft_skills": parsed.soft_skills,
                    "areas_of_interest": parsed.areas_of_interest,
                    "extra_sections": parsed.extra_sections,
                    "_raw_text": parsed.raw_text,
                }
            elif isinstance(parsed, ParsedJobDescription):
                data = {
                    "title": parsed.title,
                    "company": parsed.company,
                    "location": parsed.location,
                    "employment_type": parsed.employment_type,
                    "experience_level": parsed.experience_level,
                    "experience_years_min": parsed.experience_years_min,
                    "experience_years_max": parsed.experience_years_max,
                    "salary_min": parsed.salary_min,
                    "salary_max": parsed.salary_max,
                    "salary_currency": parsed.salary_currency,
                    "required_skills": parsed.required_skills,
                    "preferred_skills": parsed.preferred_skills,
                    "responsibilities": parsed.responsibilities,
                    "qualifications": parsed.qualifications,
                    "_raw_text": parsed.raw_text,
                }
            else:
                return

            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False)

            logger.debug(f"Saved to cache: {cache_key}")

        except Exception as e:
            logger.warning(f"Failed to save to cache: {e}")

    def clear_cache(self) -> int:
        """Clear all cached parsed documents. Returns number of files deleted."""
        cache_dir = self._get_cache_dir()
        count = 0
        for cache_file in cache_dir.glob("*.json"):
            try:
                cache_file.unlink()
                count += 1
            except Exception:
                pass
        logger.info(f"Cleared {count} cached documents")
        return count


# Global processor instance (lazy loaded)
_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get or create the document processor instance."""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
